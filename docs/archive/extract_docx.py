#!/usr/bin/env python3
import sys
import os

try:
    from docx import Document
except ImportError:
    print("Installing python-docx...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    from docx import Document

# Define document paths
docs = [
    "ProjectInfor/Sources/Project Overview.docx",
    "ProjectInfor/Sources/Project Proposal - Jerry Lelumai.docx",
    "ProjectInfor/Sources/System Design Specifications.docx",
    "ProjectInfor/Sources/Backgroud-Audit Process.docx",
    "ProjectInfor/Sources/To my project implementation.docx"
]

# Change to the correct directory
os.chdir("/home/jerry99/iaps")

# Extract content from each document
extracted_content = {}

for doc_path in docs:
    try:
        doc = Document(doc_path)
        print(f"\n{'='*80}")
        print(f"FILE: {doc_path}")
        print(f"{'='*80}\n")
        
        content = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                print(paragraph.text)
                content.append(paragraph.text)
        
        # Extract table content if any
        if doc.tables:
            print("\n--- TABLES ---")
            for i, table in enumerate(doc.tables):
                print(f"\nTable {i+1}:")
                for row in table.rows:
                    cells_content = [cell.text for cell in row.cells]
                    row_text = " | ".join(cells_content)
                    print(row_text)
                    content.append(row_text)
        
        extracted_content[doc_path] = content
    except Exception as e:
        print(f"Error reading {doc_path}: {e}")
        import traceback
        traceback.print_exc()

print(f"\n{'='*80}")
print("EXTRACTION COMPLETE")
print(f"{'='*80}\n")
